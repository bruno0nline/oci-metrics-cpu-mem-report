import os
import csv
from datetime import datetime

from docx import Document
from docx.shared import Pt

homedir = os.path.expanduser("~")
DAYS = int(os.getenv("METRICS_DAYS", "30"))

CSV_PATH = os.path.join(homedir, f"Relatorio_CPU_Memoria_media_{DAYS}d_multi_region.csv")
DOCX_PATH = os.path.join(homedir, f"Relatorio_FinOps_CPU_Mem_{DAYS}d_multi_region.docx")

# Valores aproximados por hora em USD (ajustáveis)
OCPU_PRICE_HOUR = 0.05
MEM_GB_PRICE_HOUR = 0.003
HOURS_MONTH = 730


def to_float(value):
    try:
        if value is None or value == "":
            return None
        return float(value)
    except Exception:
        return None


def estimate_monthly_cost(ocpus, mem_gb):
    ocpus = ocpus or 0
    mem_gb = mem_gb or 0
    hourly = ocpus * OCPU_PRICE_HOUR + mem_gb * MEM_GB_PRICE_HOUR
    return hourly * HOURS_MONTH


def format_money_usd(v):
    return f"US$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def load_rows():
    rows = []
    if not os.path.exists(CSV_PATH):
        print(f"CSV não encontrado: {CSV_PATH}")
        return rows

    with open(CSV_PATH, newline="") as f:
        reader = csv.DictReader(f)
        for r in reader:
            rows.append(r)
    return rows


def build_downsize_text(row):
    inst = row["instance_name"]
    shape = row["shape"]
    region = row["region"]
    comp = row["compartment"]

    cpu_mean = to_float(row["cpu_mean_percent"]) or 0
    mem_mean = to_float(row["mem_mean_percent"]) or 0
    ocpus = to_float(row["ocpus"]) or 0
    mem_gb = to_float(row["memory_gb"]) or 0

    fator = 0.5
    if cpu_mean < 5 and mem_mean < 40:
        fator = 0.25

    new_ocpus = max(1, ocpus * fator)
    new_mem = max(1, mem_gb * fator)

    current_cost = estimate_monthly_cost(ocpus, mem_gb)
    new_cost = estimate_monthly_cost(new_ocpus, new_mem)
    savings = max(0, current_cost - new_cost)

    text = (
        f"Instância: {inst} | Região: {region} | Compartment: {comp}\n"
        f"Forma atual: {shape} | OCPUs: {ocpus} | Memória: {mem_gb} GB\n"
        f"Média CPU: {cpu_mean:.2f}% | Média Memória: {mem_mean:.2f}%\n"
        f"Sugestão: reduzir para ~{new_ocpus:.1f} OCPUs e ~{new_mem:.1f} GB de memória.\n"
        f"Economia estimada: {format_money_usd(savings)}/mês.\n"
    )
    return text, savings


def build_upscale_text(row):
    inst = row["instance_name"]
    shape = row["shape"]
    region = row["region"]
    comp = row["compartment"]

    cpu_mean = to_float(row["cpu_mean_percent"]) or 0
    mem_mean = to_float(row["mem_mean_percent"]) or 0
    cpu_p95 = to_float(row["cpu_p95_percent"]) or 0
    mem_p95 = to_float(row["mem_p95_percent"]) or 0
    ocpus = to_float(row["ocpus"]) or 0
    mem_gb = to_float(row["memory_gb"]) or 0

    fator = 2.0
    new_ocpus = ocpus * fator
    new_mem = mem_gb * fator

    current_cost = estimate_monthly_cost(ocpus, mem_gb)
    new_cost = estimate_monthly_cost(new_ocpus, new_mem)
    extra = max(0, new_cost - current_cost)

    text = (
        f"Instância: {inst} | Região: {region} | Compartment: {comp}\n"
        f"Forma atual: {shape} | OCPUs: {ocpus} | Memória: {mem_gb} GB\n"
        f"CPU média/P95: {cpu_mean:.2f}% / {cpu_p95:.2f}% | "
        f"Memória média/P95: {mem_mean:.2f}% / {mem_p95:.2f}%\n"
        f"Sugestão: avaliar aumento para ~{new_ocpus:.1f} OCPUs e ~{new_mem:.1f} GB de memória.\n"
        f"Impacto estimado: +{format_money_usd(extra)}/mês.\n"
    )
    return text, extra


def generate_report():
    rows = load_rows()
    if not rows:
        return

    doc = Document()
    title = doc.add_heading(
        "Relatório FinOps – Análise de CPU e Memória (OCI)",
        level=0
    )
    title.alignment = 0

    p = doc.add_paragraph()
    run = p.add_run(
        "Gerado automaticamente a partir das métricas do OCI Monitoring. "
        "Responsável: Bruno Mendes Augusto."
    )
    run.italic = True
    run.font.size = Pt(9)

    doc.add_paragraph(f"\nJanela de análise: últimos {DAYS} dias.")

    total_down_savings = 0.0
    total_up_extra = 0.0

    # === DOWNSIZE ===
    doc.add_heading("1. Recomendações de Redução (Downsize)", level=1)
    found = False
    for r in rows:
        if r.get("finops_recommendation", "").startswith("DOWNSIZE"):
            text, savings = build_downsize_text(r)
            doc.add_paragraph(text)
            total_down_savings += savings
            found = True

    if not found:
        doc.add_paragraph("Nenhuma instância com forte indicação de redução.")

    # === UPSCALE ===
    doc.add_heading("2. Recomendações de Aumento (Upscale)", level=1)
    found = False
    for r in rows:
        if r.get("finops_recommendation") == "UPSCALE":
            text, extra = build_upscale_text(r)
            doc.add_paragraph(text)
            total_up_extra += extra
            found = True

    if not found:
        doc.add_paragraph("Nenhuma instância com forte indicação de aumento.")

    # === RESUMO ===
    doc.add_heading("3. Resumo Financeiro Consolidado (Estimativa)", level=1)

    doc.add_paragraph(
        f"Economia potencial com reduções (Downsize): "
        f"{format_money_usd(total_down_savings)}/mês."
    )

    doc.add_paragraph(
        f"Impacto potencial com aumentos (Upscale): "
        f"+{format_money_usd(total_up_extra)}/mês."
    )

    net = total_down_savings - total_up_extra
    if net >= 0:
        doc.add_paragraph(
            f"\nEconomia líquida potencial estimada: {format_money_usd(net)}/mês."
        )
    else:
        doc.add_paragraph(
            f"\nImpacto líquido potencial estimado: +{format_money_usd(abs(net))}/mês."
        )

    doc.add_paragraph(
        "\nObservação: valores estimados com base em preços de tabela simplificados. "
        "Os valores de OCPU e memória podem variar conforme contrato."
    )

    doc.save(DOCX_PATH)
    print(f"Relatório Word gerado: {DOCX_PATH}")


if __name__ == "__main__":
    generate_report()
