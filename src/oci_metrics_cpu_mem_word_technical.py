import os
import csv
from datetime import datetime

from docx import Document
from docx.shared import Pt


homedir = os.path.expanduser("~")
DAYS = int(os.getenv("METRICS_DAYS", "30"))

CSV_PATH = os.path.join(homedir, f"Relatorio_CPU_Memoria_media_{DAYS}d_multi_region.csv")
DOCX_PATH = os.path.join(homedir, f"Relatorio_FinOps_CPU_Mem_{DAYS}d_multi_region.docx")


# Preços em BRL (R$) por hora – tabela pública Oracle (estimativa)
PRICE_MATRIX = {
    "E5": {"ocpu": 0.165336,  "mem": 0.0110224},
    "E6": {"ocpu": 0.165336,  "mem": 0.0110224},
    "E4": {"ocpu": 0.13778,   "mem": 0.0082668},
    "E3": {"ocpu": 0.13778,   "mem": 0.0082668},
    "A1": {"ocpu": 0.055112,  "mem": 0.0082668},
    "A2": {"ocpu": 0.0771568, "mem": 0.0110224},
    "X9": {"ocpu": 0.220448,  "mem": 0.0082668},
}

HOURS_MONTH = 730


def infer_family(shape: str) -> str:
    if not shape:
        return "E4"
    s = shape.upper()
    for fam in ("E5", "E6", "E4", "E3", "A1", "A2", "X9"):
        if fam in s:
            return fam
    return "E4"


def get_unit_prices(shape: str):
    fam = infer_family(shape)
    prices = PRICE_MATRIX.get(fam, PRICE_MATRIX["E4"])
    return prices["ocpu"], prices["mem"]


def to_float(value):
    try:
        if value in (None, "", "no-data", "NO-DATA"):
            return None
        return float(value)
    except Exception:
        return None


def estimate_monthly_cost_brl(ocpus, mem_gb, shape):
    ocpus = ocpus or 0
    mem_gb = mem_gb or 0
    ocpu_price_hour, mem_price_hour = get_unit_prices(shape)
    hourly = ocpus * ocpu_price_hour + mem_gb * mem_price_hour
    return hourly * HOURS_MONTH


def format_money_brl(v):
    return f"R$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def load_rows():
    rows = []
    if not os.path.exists(CSV_PATH):
        print(f"CSV não encontrado: {CSV_PATH}")
        return rows

    with open(CSV_PATH, newline="") as f:
        reader = csv.DictReader(f)
        for r in reader:
            rows.append(r)
    return rows


def build_downsize(row):
    cpu_mean = to_float(row["cpu_mean_percent"]) or 0
    mem_mean = to_float(row["mem_mean_percent"]) or 0
    ocpus = to_float(row["ocpus"]) or 0
    mem_gb = to_float(row["memory_gb"]) or 0
    shape = row["shape"]

    fator = 0.5
    if cpu_mean < 5 and mem_mean < 40:
        fator = 0.25

    new_ocpus = max(1, ocpus * fator)
    new_mem = max(1, mem_gb * fator)

    current_cost = estimate_monthly_cost_brl(ocpus, mem_gb, shape)
    new_cost = estimate_monthly_cost_brl(new_ocpus, new_mem, shape)
    savings = max(0, current_cost - new_cost)

    return savings


def build_burstable(row):
    burst_enabled = row.get("burstable_enabled", "NO")
    baseline_percent = (row.get("baseline_percent") or "").strip()

    cpu_mean = to_float(row["cpu_mean_percent"]) or 0
    ocpus = to_float(row["ocpus"]) or 0
    mem_gb = to_float(row["memory_gb"]) or 0
    shape = row["shape"]

    if burst_enabled == "YES" and baseline_percent in ("12.5%", "50%"):
        return 0

    frac = None
    if cpu_mean < 8:
        frac = 0.125
    elif cpu_mean < 35:
        frac = 0.5

    if not frac:
        return 0

    current_cost = estimate_monthly_cost_brl(ocpus, mem_gb, shape)
    new_cost = estimate_monthly_cost_brl(ocpus * frac, mem_gb, shape)
    return max(0, current_cost - new_cost)


def get_top5_finops_impact(rows):
    candidates = []

    for r in rows:
        rec = r.get("finops_recommendation", "")
        if rec.startswith("DOWNSIZE"):
            savings = build_downsize(r)
        elif rec.startswith("BURSTABLE"):
            savings = build_burstable(r)
        else:
            continue

        if savings > 0:
            candidates.append((r, savings))

    candidates.sort(key=lambda x: x[1], reverse=True)
    return candidates[:5]


def generate_report():
    rows = load_rows()
    if not rows:
        return

    doc = Document()
    doc.add_heading("Relatório FinOps – OCI (CPU, Memória e Burstable)", level=0)

    p = doc.add_paragraph()
    run = p.add_run("Autor: Bruno Mendes Augusto | Relatório gerado automaticamente.")
    run.italic = True
    run.font.size = Pt(9)

    doc.add_paragraph(f"Janela de análise: últimos {DAYS} dias.")
    doc.add_paragraph("Valores estimados em real brasileiro (BRL).")

    # ================= TOP 5 =================
    doc.add_heading("🏆 TOP 5 Oportunidades de Economia (Baixo Risco)", level=1)

    top5 = get_top5_finops_impact(rows)
    total_top5 = 0

    if top5:
        table = doc.add_table(rows=1, cols=6)
        hdr = table.rows[0].cells
        hdr[0].text = "Rank"
        hdr[1].text = "Instância"
        hdr[2].text = "Região"
        hdr[3].text = "Shape"
        hdr[4].text = "Recomendação"
        hdr[5].text = "Economia Estimada (R$/mês)"

        for idx, (r, savings) in enumerate(top5, start=1):
            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = r["instance_name"]
            row[2].text = r["region"]
            row[3].text = r["shape"]
            row[4].text = r["finops_recommendation"]
            row[5].text = format_money_brl(savings)
            total_top5 += savings

        doc.add_paragraph(f"\nEconomia potencial total do TOP 5: {format_money_brl(total_top5)}/mês.")
    else:
        doc.add_paragraph("Nenhuma oportunidade relevante identificada.")

    # ================= RESUMO =================
    doc.add_heading("📊 Resumo Executivo", level=1)
    doc.add_paragraph(
        "Recomendação estratégica: atuar mensalmente apenas sobre as TOP 5 instâncias, "
        "minimizando risco operacional e maximizando retorno financeiro."
    )

    doc.add_paragraph(
        "\nObservação: valores estimados com base em preços públicos OCI. "
        "Licenças de sistema operacional não estão incluídas."
    )

    doc.save(DOCX_PATH)
    print(f"Relatório Word gerado: {DOCX_PATH}")


if __name__ == "__main__":
    generate_report()
