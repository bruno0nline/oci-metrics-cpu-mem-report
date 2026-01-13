import os
import csv
from datetime import datetime
from docx import Document
from docx.shared import Pt

DAYS = int(os.getenv("METRICS_DAYS", "30"))
HOME = os.path.expanduser("~")

CSV_PATH = os.path.join(HOME, f"Relatorio_FinOps_CPU_MEM_{DAYS}d.csv")
DOCX_PATH = os.path.join(HOME, f"Relatorio_FinOps_Downsizes_Strong_{DAYS}d.docx")

OCPU_PRICE = 0.05
MEM_PRICE = 0.003
HOURS_MONTH = 730


def to_float(v):
    try:
        return float(v)
    except Exception:
        return 0.0


def estimate_cost(ocpus, mem):
    return (ocpus * OCPU_PRICE + mem * MEM_PRICE) * HOURS_MONTH


def format_usd(v):
    return f"US$ {v:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def main():
    rows = []
    with open(CSV_PATH, newline="", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for r in reader:
            if r["finops_recommendation"] == "DOWNSIZE-STRONG":
                rows.append(r)

    doc = Document()
    doc.add_heading("Relatório FinOps – Oportunidades de Economia", level=0)

    doc.add_paragraph(
        f"Janela de análise: últimos {DAYS} dias\n"
        f"Gerado em: {datetime.now().strftime('%d/%m/%Y %H:%M')}"
    ).italic = True

    total = 0.0

    for r in rows:
        ocpus = to_float(r["ocpus"])
        mem = to_float(r["memory_gb"])

        new_ocpus = max(1, ocpus * 0.5)
        new_mem = max(1, mem * 0.5)

        cur = estimate_cost(ocpus, mem)
        new = estimate_cost(new_ocpus, new_mem)
        savings = cur - new
        total += savings

        doc.add_paragraph(
            f"Instância: {r['instance_name']}\n"
            f"Região: {r['region']} | Compartment: {r['compartment']}\n"
            f"OCPUs: {ocpus} → {new_ocpus:.1f}\n"
            f"Memória: {mem} GB → {new_mem:.1f} GB\n"
            f"Economia estimada: {format_usd(savings)}/mês\n"
        )

    doc.add_heading("Resumo financeiro", level=1)
    doc.add_paragraph(
        f"Economia total potencial estimada: {format_usd(total)}/mês."
    )

    doc.save(DOCX_PATH)
    print(f"✅ Relatório Word gerado: {DOCX_PATH}")


if __name__ == "__main__":
    main()
